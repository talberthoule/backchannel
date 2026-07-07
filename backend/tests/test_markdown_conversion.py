import io
import unittest

from app.services.markdown_conversion import (
    MarkdownConversionError,
    convert_to_markdown,
    file_extension,
)


def _docx_bytes(heading: str, paragraph: str) -> bytes:
    import docx

    buf = io.BytesIO()
    d = docx.Document()
    d.add_heading(heading, level=1)
    d.add_paragraph(paragraph)
    d.save(buf)
    return buf.getvalue()


class FileExtensionTests(unittest.TestCase):
    def test_lowercases_and_extracts(self):
        self.assertEqual(file_extension("Report.DOCX"), ".docx")
        self.assertEqual(file_extension("notes.md"), ".md")
        self.assertEqual(file_extension("noext"), "")
        self.assertEqual(file_extension(""), "")


class ConvertToMarkdownTests(unittest.TestCase):
    def test_markdown_and_text_pass_through_unconverted(self):
        body = "# Heading\n\nSome *markdown* content."
        self.assertEqual(convert_to_markdown(body.encode(), "notes.md"), body)
        self.assertEqual(convert_to_markdown(b"plain text\n", "notes.txt"), "plain text")

    def test_docx_converts_to_markdown(self):
        content = _docx_bytes("Service Catalog", "Managed firewall service.")
        markdown = convert_to_markdown(content, "catalog.docx")
        self.assertIn("# Service Catalog", markdown)
        self.assertIn("Managed firewall service.", markdown)

    def test_csv_converts_to_markdown_table(self):
        content = b"title,body\nService A,Does A things\n"
        markdown = convert_to_markdown(content, "records.csv")
        self.assertIn("Service A", markdown)
        self.assertIn("|", markdown)

    def test_unsupported_extension_raises(self):
        with self.assertRaises(MarkdownConversionError):
            convert_to_markdown(b"binary", "audio.wav")

    def test_empty_docx_raises(self):
        content = _docx_bytes("", "")
        with self.assertRaises(MarkdownConversionError):
            convert_to_markdown(content, "empty.docx")


if __name__ == "__main__":
    unittest.main()
